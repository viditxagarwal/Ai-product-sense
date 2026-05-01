"""File upload, parsing, context assembly, and artifact persistence helpers."""

import csv
import io
import json
import os
import re
from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import HTTPException, UploadFile
from openpyxl import load_workbook

from app.config import UPLOAD_DIR
from app.database import supabase


TEXT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/json",
}


def _safe_filename(filename: str) -> str:
    stem = Path(filename).stem or "file"
    suffix = Path(filename).suffix.lower()
    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._") or "file"
    return f"{safe_stem[:80]}{suffix[:16]}"


def _thread_upload_dir(thread_id: str) -> Path:
    path = Path(UPLOAD_DIR) / str(thread_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _public_upload_url(thread_id: str, stored_name: str) -> str:
    return f"/uploads/{thread_id}/{stored_name}"


def _path_from_file_url(file_url: str) -> Path | None:
    if not file_url.startswith("/uploads/"):
        return None
    relative = file_url.removeprefix("/uploads/").lstrip("/")
    path = (Path(UPLOAD_DIR) / relative).resolve()
    upload_root = Path(UPLOAD_DIR).resolve()
    if upload_root not in path.parents:
        return None
    return path


async def save_thread_upload(user_id: str, thread_id: str, upload: UploadFile) -> dict:
    """Persist an uploaded file locally and create the thread_files record."""
    thread = (
        supabase.table("threads")
        .select("id")
        .eq("id", str(thread_id))
        .eq("user_id", str(user_id))
        .single()
        .execute()
    )
    if not thread.data:
        raise HTTPException(status_code=404, detail="Thread not found")

    original_name = upload.filename or "upload.bin"
    safe_name = _safe_filename(original_name)
    stored_name = f"{uuid4().hex}_{safe_name}"
    file_path = _thread_upload_dir(thread_id) / stored_name

    content = await upload.read()
    file_path.write_bytes(content)

    file_type = upload.content_type or _guess_mime(original_name)
    payload = {
        "thread_id": str(thread_id),
        "file_name": original_name,
        "file_url": _public_upload_url(thread_id, stored_name),
        "file_type": file_type,
        "file_size_bytes": len(content),
        "source": "user_upload",
    }
    resp = supabase.table("thread_files").insert(payload).execute()
    record = resp.data[0]

    supabase.table("file_versions").insert({
        "file_id": record["id"],
        "version_number": 1,
        "file_url": record["file_url"],
        "operation_type": "creation",
        "change_summary": {"source": "user_upload", "bytes": len(content)},
        "created_by": "user",
    }).execute()

    return record


def build_thread_file_context(thread_id: str, max_chars: int = 12000) -> dict[str, Any]:
    """Parse user-uploaded files for a thread into prompt-ready context."""
    resp = (
        supabase.table("thread_files")
        .select("*")
        .eq("thread_id", str(thread_id))
        .eq("source", "user_upload")
        .order("created_at", desc=False)
        .execute()
    )
    files = resp.data or []
    if not files:
        return {"text": "", "files": [], "total_chars": 0}

    remaining = max(max_chars, 0)
    sections: list[str] = []
    summaries: list[dict[str, Any]] = []

    for file in files:
        parsed = extract_file_text(file, max_chars=remaining)
        summaries.append({
            "id": file.get("id"),
            "name": file.get("file_name"),
            "type": file.get("file_type"),
            "size_bytes": file.get("file_size_bytes"),
            "chars": len(parsed["text"]),
            "status": parsed["status"],
            "note": parsed.get("note"),
        })
        if parsed["text"] and remaining > 0:
            sections.append(
                f"### File: {file.get('file_name')}\n"
                f"Type: {file.get('file_type')}\n"
                f"Parsed content:\n{parsed['text']}"
            )
            remaining -= len(parsed["text"])

    text = "\n\n".join(sections)
    return {"text": text, "files": summaries, "total_chars": len(text)}


def extract_file_text(file: dict, max_chars: int = 12000) -> dict[str, str]:
    path = _path_from_file_url(file.get("file_url", ""))
    file_name = file.get("file_name", "")
    file_type = (file.get("file_type") or "").lower()

    if not path or not path.exists():
        return {"status": "missing", "text": "", "note": "File bytes are not available on disk."}

    try:
        if file_type.startswith("image/"):
            return {
                "status": "metadata_only",
                "text": "",
                "note": "Image bytes are stored for preview, but OCR/multimodal prompt injection is not enabled.",
            }
        if file_type == "application/pdf" or file_name.lower().endswith(".pdf"):
            return _extract_pdf(path, max_chars)
        if file_name.lower().endswith(".docx") or "wordprocessingml" in file_type:
            return _extract_docx(path, max_chars)
        if file_name.lower().endswith((".xlsx", ".xlsm")) or "spreadsheet" in file_type:
            return _extract_xlsx(path, max_chars)
        if file_type in TEXT_TYPES or file_name.lower().endswith((".txt", ".md", ".csv", ".json", ".py", ".ts", ".js", ".sql")):
            raw = path.read_bytes()
            text = raw.decode("utf-8", errors="replace")
            if file_type == "application/json" or file_name.lower().endswith(".json"):
                text = _format_json(text)
            if file_type == "text/csv" or file_name.lower().endswith(".csv"):
                text = _format_csv(text, max_rows=120)
            return {"status": "parsed", "text": text[:max_chars], "note": ""}
        return {"status": "unsupported", "text": "", "note": f"No parser for {file_type or 'unknown type'}."}
    except Exception as exc:
        return {"status": "error", "text": "", "note": str(exc)}


def persist_artifacts_from_output(
    thread_id: str,
    output_text: str,
    trigger_step_id: str | None = None,
) -> list[dict[str, str]]:
    """Persist fenced artifact blocks from an LLM response as thread files."""
    created: list[dict[str, str]] = []
    for match in re.finditer(r"```artifact(?P<attrs>[^\n`]*)\n(?P<body>.*?)```", output_text, re.DOTALL | re.IGNORECASE):
        attrs = match.group("attrs") or ""
        body = match.group("body").strip("\n")
        filename = _extract_attr(attrs, "filename") or f"artifact-{len(created) + 1}.md"
        mime = _extract_attr(attrs, "type") or _guess_mime(filename)
        record = persist_generated_artifact(thread_id, filename, body, mime, trigger_step_id)
        created.append({
            "file_id": record["id"],
            "file_name": record["file_name"],
            "file_type": record["file_type"],
        })
    return created


def persist_generated_artifact(
    thread_id: str,
    filename: str,
    content: str,
    mime_type: str,
    trigger_step_id: str | None = None,
) -> dict:
    safe_name = _safe_filename(filename)
    stored_name = f"{uuid4().hex}_{safe_name}"
    file_path = _thread_upload_dir(thread_id) / stored_name
    data = content.encode("utf-8")
    file_path.write_bytes(data)

    record = supabase.table("thread_files").insert({
        "thread_id": str(thread_id),
        "file_name": filename,
        "file_url": _public_upload_url(thread_id, stored_name),
        "file_type": mime_type,
        "file_size_bytes": len(data),
        "source": "ai_generated",
    }).execute().data[0]

    version_payload: dict[str, Any] = {
        "file_id": record["id"],
        "version_number": 1,
        "file_url": record["file_url"],
        "operation_type": "creation",
        "change_summary": {"source": "llm_artifact_block"},
        "created_by": "ai",
    }
    if trigger_step_id:
        version_payload["trigger_step_id"] = trigger_step_id
    supabase.table("file_versions").insert(version_payload).execute()
    return record


def persist_generated_artifact_bytes(
    thread_id: str,
    filename: str,
    content: bytes,
    mime_type: str,
    trigger_step_id: str | None = None,
    change_summary: dict[str, Any] | None = None,
) -> dict:
    safe_name = _safe_filename(filename)
    stored_name = f"{uuid4().hex}_{safe_name}"
    file_path = _thread_upload_dir(thread_id) / stored_name
    file_path.write_bytes(content)

    record = supabase.table("thread_files").insert({
        "thread_id": str(thread_id),
        "file_name": filename,
        "file_url": _public_upload_url(thread_id, stored_name),
        "file_type": mime_type,
        "file_size_bytes": len(content),
        "source": "ai_generated",
    }).execute().data[0]

    version_payload: dict[str, Any] = {
        "file_id": record["id"],
        "version_number": 1,
        "file_url": record["file_url"],
        "operation_type": "creation",
        "change_summary": change_summary or {"source": "file_writer"},
        "created_by": "ai",
    }
    if trigger_step_id:
        version_payload["trigger_step_id"] = trigger_step_id
    supabase.table("file_versions").insert(version_payload).execute()
    return record


def _extract_pdf(path: Path, max_chars: int) -> dict[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return {
            "status": "unsupported",
            "text": "",
            "note": "PDF parsing requires the pypdf package.",
        }

    reader = PdfReader(str(path))
    chunks = []
    for page in reader.pages[:40]:
        chunks.append(page.extract_text() or "")
        if sum(len(chunk) for chunk in chunks) >= max_chars:
            break
    return {"status": "parsed", "text": "\n\n".join(chunks)[:max_chars], "note": ""}


def _extract_docx(path: Path, max_chars: int) -> dict[str, str]:
    try:
        from docx import Document
    except ImportError:
        return {
            "status": "unsupported",
            "text": "",
            "note": "DOCX parsing requires the python-docx package.",
        }

    doc = Document(str(path))
    lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    return {"status": "parsed", "text": "\n".join(lines)[:max_chars], "note": ""}


def _extract_xlsx(path: Path, max_chars: int) -> dict[str, str]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"Sheet: {sheet.title}")
        for idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if idx > 120:
                lines.append("... rows truncated ...")
                break
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                lines.append(", ".join(values))
            if sum(len(line) for line in lines) >= max_chars:
                return {"status": "parsed", "text": "\n".join(lines)[:max_chars], "note": ""}
    return {"status": "parsed", "text": "\n".join(lines)[:max_chars], "note": ""}


def _format_json(text: str) -> str:
    try:
        return json.dumps(json.loads(text), indent=2)
    except json.JSONDecodeError:
        return text


def _format_csv(text: str, max_rows: int) -> str:
    rows = list(csv.reader(io.StringIO(text)))
    return "\n".join(", ".join(row) for row in rows[:max_rows])


def _extract_attr(attrs: str, key: str) -> str | None:
    match = re.search(rf'{key}=["\']([^"\']+)["\']', attrs)
    return match.group(1) if match else None


def _guess_mime(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".md"):
        return "text/markdown"
    if lower.endswith(".csv"):
        return "text/csv"
    if lower.endswith(".json"):
        return "application/json"
    if lower.endswith(".pdf"):
        return "application/pdf"
    if lower.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if lower.endswith(".xlsx"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if lower.endswith((".png", ".jpg", ".jpeg")):
        return "image/png" if lower.endswith(".png") else "image/jpeg"
    return "text/plain"
